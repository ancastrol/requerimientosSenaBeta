import streamlit as st
import pandas as pd
from PyPDF2 import PdfMerger
import io
from io import BytesIO
from datetime import datetime, timedelta
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from dateutil.relativedelta import relativedelta
from docx import Document
import os
from zipfile import ZipFile
import tempfile
from tqdm import tqdm
from docx2pdf import convert
import pythoncom
import win32com.client
from win32com.client import constants
import time
import msvcrt

# Leer excel prueba 
df = pd.read_excel(r'Sources/prototiposSeguimiento.xlsx')

# Inicialización del estado de sesión
if 'vista_actual' not in st.session_state:
    st.session_state.vista_actual = 'inicio'

# Recibe los documentos excel y word y regresa un zip con los documentos word llenos
def procesar_texto(texto, aprendiz, columnas):

    texto_nuevo = texto
    
    # Procesar marcadores especiales de nivel
    if "{TECNICO}" in texto_nuevo:
        if aprendiz['Nivel'] == "Tecnico":
            texto_nuevo = texto_nuevo.replace("{TECNICO}", "X")
        else:
            texto_nuevo = texto_nuevo.replace("{TECNICO}", "")
    
    if "{TECNOLOGO}" in texto_nuevo:
        if aprendiz['Nivel'] == "Tecnologo":
            texto_nuevo = texto_nuevo.replace("{TECNOLOGO}", "X")
        else:
            texto_nuevo = texto_nuevo.replace("{TECNOLOGO}", "")

    # Procesar marcadores especiales de etapa productiva
    if "{CA}" in texto_nuevo:
        if aprendiz['Alternativa(Etapa Productiva)'] == "CA":
            texto_nuevo = texto_nuevo.replace("{CA}", "X")
        else:
            texto_nuevo = texto_nuevo.replace("{CA}", "")
    if "{VL}" in texto_nuevo:
        if aprendiz['Alternativa(Etapa Productiva)'] == "VL":
            texto_nuevo = texto_nuevo.replace("{VL}", "X")
        else:
            texto_nuevo = texto_nuevo.replace("{VL}", "")
    if "{P}" in texto_nuevo:
        if aprendiz['Alternativa(Etapa Productiva)'] == "PP":
            texto_nuevo = texto_nuevo.replace("{P}", "X")
        else:
            texto_nuevo = texto_nuevo.replace("{P}", "")
    if "{PA}" in texto_nuevo:
        if aprendiz['Alternativa(Etapa Productiva)'] == "PA":
            texto_nuevo = texto_nuevo.replace("{PA}", "X")
        else:
            texto_nuevo = texto_nuevo.replace("{PA}", "")
    
    # Procesar el resto de marcadores
    for columna in columnas:
        marcador = '{' + columna.upper() + '}'
        if marcador in texto_nuevo:
            texto_nuevo = texto_nuevo.replace(marcador, str(aprendiz[columna]))
    
    return texto_nuevo

def procesar_parrafos(container, aprendiz, columnas):
    for paragraph in container.paragraphs:
        for run in paragraph.runs:
            run.text = procesar_texto(run.text, aprendiz, columnas)

def procesar_tablas(doc, aprendiz, columnas):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                procesar_parrafos(cell, aprendiz, columnas)

def procesar_documentos(df, plantilla_word):

    zip_archivo = io.BytesIO()
    documentos_exitosos = 0
    errores = []
    
    with ZipFile(zip_archivo, 'w') as zip_file:
        plantilla_temp = io.BytesIO(plantilla_word)
        columnas = df.columns
        
        for index, aprendiz in df.iterrows():
            try:
                # Cargar nueva copia de la plantilla
                doc = Document(plantilla_temp)
                plantilla_temp.seek(0)
                
                # Procesar el documento
                procesar_parrafos(doc, aprendiz, columnas)
                procesar_tablas(doc, aprendiz, columnas)
                
                # Guardar el documento procesado
                doc_temp = io.BytesIO()
                doc.save(doc_temp)
                doc_temp.seek(0)
                
                # Generar nombre del archivo
                nombre_archivo = f"{aprendiz['Nombre']}_{aprendiz['Apellidos']}_acta_de_inicio.docx"
                
                # Añadir al ZIP
                zip_file.writestr(nombre_archivo, doc_temp.getvalue())
                documentos_exitosos += 1
                
            except Exception as e:
                errores.append(f"Error procesando aprendiz {aprendiz['Nombre']} {aprendiz['Apellidos']}: {str(e)}")
    
    zip_archivo.seek(0)
    return zip_archivo, documentos_exitosos, errores

def desercion_12_meses(df):
    # Se obtiene la fecha actual
    fecha_actual = pd.Timestamp.now()
    
    # Se hace que la columna de fechas esté en formato datetime
    df['Inicio Ficha'] = pd.to_datetime(df['Inicio Ficha'], errors='coerce')
    print(df['Inicio Ficha'].dtype)
    
    # Se realiza una lista de aprendices que no han elegido etapa productiva en el tiempo establecido
    aprendices_descercion = df[(df['Inicio Ficha'] < (fecha_actual - relativedelta(months=12))) & (df['Alternativa(Etapa Productiva)'].str.upper() == 'NO')]
    
    # Iterar sobre las filas del DataFrame de deserción
    for index, aprendiz in aprendices_descercion.iterrows():
        enviar_correo_desercion(
            'Proceso de descerción', 
            f'El aprendiz {aprendiz["Aprendiz"]} no ha elegido una alternativa de etapa productiva en el tiempo establecido', 
            aprendiz['instructor seguimiento'], 
            aprendiz['Correo Aprendiz'])

# Función para buscar un aprendiz en el DataFrame
def buscar_aprendiz(df, documento):
    # Convertir el número de documento a string y eliminar espacios
    documento = str(documento).strip()
    
    # Convertir la columna de documento a string y eliminar espacios
    df['NUMERO DE DOCUMENTO'] = df['NUMERO DE DOCUMENTO'].astype(str).str.strip()
    
    # Buscar el aprendiz
    aprendiz = df[df['NUMERO DE DOCUMENTO'] == documento]
    
    if len(aprendiz) == 0:
        print(f"No se encontró ningún aprendiz con el documento {documento}")
        return None
    
    if len(aprendiz) > 1:
        print(f"Advertencia: Se encontraron {len(aprendiz)} aprendices con la identificacion No. {documento}")
    
    return aprendiz.iloc[0]

def convertir_a_pdf(bytes_word):
    pythoncom.CoInitialize()
    word = None
    temp_dir = None
    
    try:
        # Crear directorio temporal
        temp_dir = os.path.join(os.environ['TEMP'], f'word_to_pdf_{int(time.time())}')
        os.makedirs(temp_dir, exist_ok=True)
        
        # Nombres de archivos
        temp_docx = os.path.join(temp_dir, 'documento.docx')
        temp_pdf = os.path.join(temp_dir, 'documento.pdf')
        
        # Escribir archivo Word
        with open(temp_docx, 'wb') as f:
            f.write(bytes_word.getvalue())
            f.flush()
            os.fsync(f.fileno())
        
        # Verificar archivo Word
        if not os.path.exists(temp_docx):
            raise Exception("Error al crear archivo Word temporal")
            
        # Esperar a que el archivo esté disponible
        time.sleep(1)
        
        # Inicializar Word con configuración específica
        word = win32com.client.DispatchEx('Word.Application')
        word.Visible = False
        word.DisplayAlerts = False
        
        try:
            # Abrir documento
            doc = word.Documents.Open(os.path.abspath(temp_docx))
            
            # Configurar opciones de PDF
            doc.SaveAs2(
                FileName=temp_pdf,
                FileFormat=17,  # wdFormatPDF
                CompatibilityMode=15
            )
            
            # Cerrar documento
            doc.Close(SaveChanges=False)
            
            # Esperar a que el PDF esté listo
            time.sleep(2)
            
            if not os.path.exists(temp_pdf):
                raise Exception("Error al generar PDF")
                
            # Verificar tamaño del PDF
            if os.path.getsize(temp_pdf) < 100:  # menos de 100 bytes probablemente es corrupto
                raise Exception("PDF generado está vacío o corrupto")
                
            # Leer PDF
            pdf_bytes = BytesIO()
            with open(temp_pdf, 'rb') as f:
                pdf_content = f.read()
                if not pdf_content:
                    raise Exception("No se pudo leer el contenido del PDF")
                pdf_bytes.write(pdf_content)
            
            pdf_bytes.seek(0)
            return pdf_bytes
            
        finally:
            if doc:
                try:
                    doc.Close(SaveChanges=False)
                except:
                    pass
            
    except Exception as e:
        raise Exception(f"Error en la conversión: {str(e)}")
        
    finally:
        # Cerrar Word
        if word:
            try:
                word.Quit()
            except:
                pass
            
        # Limpiar archivos temporales
        try:
            if os.path.exists(temp_docx):
                os.unlink(temp_docx)
            if os.path.exists(temp_pdf):
                os.unlink(temp_pdf)
            if temp_dir and os.path.exists(temp_dir):
                os.rmdir(temp_dir)
        except:
            pass
            
        pythoncom.CoUninitialize()

def procesar_documento_individual(documento_aprendiz, df, plantilla_word):
    try:
        # Verificar entrada
        if not plantilla_word:
            raise Exception("Plantilla Word vacía")
            
        # Buscar aprendiz
        aprendiz = buscar_aprendiz(df, documento_aprendiz)
        if aprendiz is None:
            st.error(f"No se encontró el aprendiz con documento de indentidad: {documento_aprendiz}")
            return []

        # Cargar plantilla
        doc = Document(BytesIO(plantilla_word))
        
        # Procesar documento
        columnas = df.columns
        procesar_parrafos(doc, aprendiz, columnas)
        procesar_tablas(doc, aprendiz, columnas)

        # Guardar documento procesado
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)

        # Convertir a PDF
        pdf_buffer = convertir_a_pdf(doc_bytes)
        
        if not pdf_buffer:
            raise Exception("Fallo en la conversión a PDF")
            
        return [pdf_buffer]

    except Exception as e:
        st.error(f"Error: {str(e)}")
        return []

# Función para cambiar la vista actual de la aplicación
def cambiar_vista(nueva_vista):
    st.session_state.vista_actual = nueva_vista

# Función para mostrar la vista de inicio  
def mostrar_inicio():
    st.title('Elija su rol para continuar')
    st.subheader('Seleccione una de las siguientes opciones para continuar con las herramientas indicadas para su rol.')
    st.write('')
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button('Instructor', key='btn_instructor'):
            cambiar_vista('instructor')
            st.rerun()
            
    with col2:
        if st.button('Aprendíz', key='btn_aprendiz'):
            cambiar_vista('aprendiz')
            st.rerun()

# Función para mostrar la vista de instructor
def mostrar_instructor():
    if st.button('⬅️ Volver atrás', key='volver_instructor'):
        cambiar_vista('inicio')
        st.rerun()
        
    # Crear la barra lateral
    st.sidebar.title("Menú desplegable")
    opcion = st.sidebar.selectbox(
        'Elige una opción:',
        ('Pantalla inicial', 'Consolidado PDF', 'Cruce de correspondencia', 'Proceso de descercion')
    )
    
    if opcion == 'Pantalla inicial':

        col1, col2 = st.columns(2)

        with col1:
            st.title('Herramientas de Desarrollo Etapa Productiva')

        with col2:
            # Mostrar la imagen en la ventana
            st.image("picture103.jpg", width=200)

        # Contenido de la pagina
        st.write('Este aplicativo busca facilitar multiples tareas de los instructores con respecto al manejo de sus aprendices que estan terminando la etapa productiva y estan en curso de certificarse. Si desea ver las funcionalidades disponibles se encuentra en la barra lateral a la izquierda de la pantalla.')

    elif opcion == 'Consolidado PDF':
        mostrar_aprendiz()

    elif opcion == 'Cruce de correspondencia':

        st.title("🎓 Generador de Documentos para Aprendices")
        
        # Agregar información de uso
        with st.expander("ℹ️ Instrucciones de uso"):
            st.markdown("""
            1. **Preparación del Excel:**
            * Asegúrate de que tu Excel tenga encabezados claros
            * Cada columna debe corresponder a un marcador en la plantilla Word
            
            2. **Preparación de la Plantilla Word:**
            * Usa marcadores entre llaves, ejemplo: {NOMBRE}, {APELLIDO}
            * Los marcadores deben coincidir con los nombres de las columnas del Excel
            
            3. **Proceso:**
            * Sube tu archivo Excel con los datos
            * Sube tu plantilla Word
            * Haz clic en 'Generar Documentos'
            * Descarga el archivo ZIP con todos los documentos generados
            """)
        
        # Columnas para la carga de archivos
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("📊 Archivo Excel")
            excel_file = st.file_uploader("Sube tu archivo Excel", type=['xlsx', 'xls'])
            
            if excel_file is not None:
                try:
                    dfSubido = pd.read_excel(excel_file)
                    st.success(f"✅ Excel cargado exitosamente - {len(dfSubido)} registros encontrados")
                    
                    # Mostrar vista previa de los datos
                    with st.expander("👀 Vista previa de datos"):
                        st.dataframe(df.head())
                    
                    # Mostrar los marcadores disponibles
                    st.info("🔍 Marcadores disponibles:")
                    marcadores = [f"{{{col.upper()}}}" for col in dfSubido.columns]
                    st.code(", ".join(marcadores))
                except Exception as e:
                    st.error(f"❌ Error al leer el archivo Excel: {str(e)}")
                    df = None
        
        with col2:
            st.subheader("📄 Plantilla Word")
            word_file = st.file_uploader("Sube tu plantilla Word", type=['docx'])
            
            if word_file is not None:
                st.success("✅ Plantilla Word cargada exitosamente")
        
        # Botón para generar documentos
        if st.button("🚀 Generar Documentos", disabled=(excel_file is None or word_file is None)):
            if excel_file is not None and word_file is not None:
                with st.spinner("Generando documentos..."):
                    try:
                        # Procesar documentos
                        zip_archivo, docs_exitosos, errores = procesar_documentos(df, word_file.getvalue())
                        
                        # Mostrar resultados
                        st.success(f"✅ Proceso completado - {docs_exitosos} documentos generados")
                        
                        # Si hay errores, mostrarlos
                        if errores:
                            with st.expander("⚠️ Errores encontrados"):
                                for error in errores:
                                    st.error(error)
                        
                        # Botón de descarga
                        st.download_button(
                            label="📥 Descargar Documentos (ZIP)",
                            data=zip_archivo.getvalue(),
                            file_name="documentos_generados.zip",
                            mime="application/zip"
                        )
                        
                    except Exception as e:
                        st.error(f"❌ Error durante la generación: {str(e)}")
    elif opcion == 'Proceso de descercion':
        st.title("Verificar aprendices que deban iniciar proceso de descerción")
        
        # Columnas para la carga de archivos
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("📊 Archivo Excel")
            excel_file = st.file_uploader("Sube tu archivo Excel", type=['xlsx', 'xls'])
            
            if excel_file is not None:
                try:
                    df = pd.read_excel(excel_file)
                    st.success(f"Excel cargado exitosamente - {len(df)} registros encontrados")
                    
                    # Boton que ejecuta la funcion de descercion
                    if st.button("🚀 Verificar descerción"):
                        if excel_file is not None:
                            with st.spinner("Verificando descerción..."):
                                try:
                                    desercion_12_meses(df)
                                    st.success(f"✅ Proceso de descerción completado")
                                except Exception as e:
                                    st.error(f"❌ Error durante el proceso de descerción: {str(e)}")

                except Exception as e:
                    st.error(f"Error al leer el archivo Excel: {str(e)}")
                    df = None
        

# Función para mostrar la vista de aprendiz
def mostrar_aprendiz():

    # Botón para volver a la vista de inicio
    if st.button('⬅️ Volver atrás', key='volver_aprendiz'):
        cambiar_vista('inicio')
        st.rerun()

    # Titulo de la página
    st.title('LLene este formulario para subir los documentos requeridos para la finalización de la etapa productiva')

    # Formulario para almacenar los datos faltantes en el excel
    st.subheader('Formulario de subida de documentos')

    # Campos de entrada
    col1, col2 = st.columns(2)
    with col1:
        documento = st.text_input('Documento de identidad:')
        nombre = st.text_input('Nombre del aprendiz:')
    with col2:
        ficha = st.text_input('Numero de ficha:')

    # Se dirige a la vista del consolidado pdf
    if st.button('Consolidado PDF', key='btn_consolidado_pdf'):
        st.session_state.nombre = nombre
        st.session_state.documento = documento
        st.session_state.ficha = ficha
        cambiar_vista('formulario')
        st.rerun()
        
        

# Funcion para mostrar la vista del formulario de consolidado PDF
def mostrar_formulario():
    # Botón para volver a la vista de inicio
    if st.button('⬅️ Volver atrás', key='volver_aprendiz'):
        cambiar_vista('inicio')
        st.rerun()

    st.title('Consolidado PDF')
    mostrar_formulario_pdf()


# Función para enviar correo al instructor con copia al aprendiz y adjunto
def enviar_correo_instructor(asunto, cuerpo, archivo_pdf, nombre_archivo):

    # Obtener el destinatarios del correo electrónico
    destinatario = df['instructor seguimiento'].iloc[0]
    destinatario2 = df['Correo Aprendiz'].iloc[0]

    # Obtener los detalles del correo electrónico
    msg = MIMEMultipart()
    msg['From'] = 'astroc2208@gmail.com'
    msg['To'] = str(destinatario)
    msg['Cc'] = str(destinatario2)
    msg['Subject'] = asunto
    msg.attach(MIMEText(cuerpo, 'plain'))

    # Adjuntar el archivo PDF
    adjunto = MIMEApplication(archivo_pdf.getvalue(), _subtype="pdf")
    adjunto.add_header('Content-Disposition', 'attachment', filename = nombre_archivo)
    msg.attach(adjunto)

    # Conexión al servidor SMTP
    smtp = smtplib.SMTP('smtp.gmail.com', 587)
    smtp.starttls()

    # Autenticación con tu correo y contraseña de aplicación
    smtp.login('astroc2208@gmail.com', 'jsgm gpyz gakh ywop')

    # Envío del correo
    smtp.send_message(msg)

    # Cierre de la conexión SMTP
    smtp.quit()

# Función para enviar correo al instructor con copia al aprendiz y adjunto
def enviar_correo_desercion(asunto, cuerpo, destinatario, destinatario2):

    # Obtener los detalles del correo electrónico
    msg = MIMEMultipart()
    msg['From'] = 'astroc2208@gmail.com'
    msg['To'] = str(destinatario)
    msg['Cc'] = str(destinatario2)
    msg['Subject'] = asunto
    msg.attach(MIMEText(cuerpo, 'plain'))

    # Conexión al servidor SMTP
    smtp = smtplib.SMTP('smtp.gmail.com', 587)
    smtp.starttls()

    # Autenticación con tu correo y contraseña de aplicación
    smtp.login('astroc2208@gmail.com', 'jsgm gpyz gakh ywop')

    # Envío del correo
    smtp.send_message(msg)

    # Cierre de la conexión SMTP
    smtp.quit()

# Función para verificar si la ficha se encuentra y si corresponde a un tecnólogo
def verificar_ficha_tecnologo(df, numero_ficha):
    # Verificar si el DataFrame no esta vacio
    if df is None:
        return False
    
    # Convertir la columna de ficha a string para comparación
    df['Ficha'] = df['Ficha'].astype(str)
    
    # Buscar la ficha en el DataFrame
    ficha_encontrada = df[df['Ficha'] == str(numero_ficha)]
    
    if len(ficha_encontrada) > 0:
        # Devolver el valor de Es_Tecnologo
        return ficha_encontrada['Nivel'].iloc[0]
    else:
        st.warning(f"Ficha {numero_ficha} no encontrada en el registro")
        return False
    
# Funcion para unir PDFs, tiene en cuenta que el TyT de haberlo va antes del ultimo documento
def unir_pdfs_con_orden(archivos_pdf, es_tecnologo):
    merger = PdfMerger()
    
    # Se agregan los PDFs
    otros_docs = []
    
    for archivo in archivos_pdf:
        if archivo is not None:
            # Si es el documento de tecnólogo, se guarda aparte
            otros_docs.append(archivo)

    # Se crea el PDF final
    output = io.BytesIO()
    merger.write(output)
    output.seek(0)
    return output

def mostrar_formulario_pdf():
    try:
        # Cargar datos
        fichaPdf = st.session_state.ficha
        nombrePdf = st.session_state.nombre
        documentoPdf = st.session_state.documento
        
        st.write('Suba los archivos PDF en el campo correspondiente, al oprimir el botón "Consolidar PDFs" se unirán los archivos en un solo PDF y se enviará al instructor.')

        # Procesar documento con manejo de bytes
        try:
            with open('Sources/Formato-entrega-documentacion-V7.docx', 'rb') as plantilla:
                plantilla_bytes = plantilla.read()
                formato = procesar_documento_individual(documentoPdf, df, plantilla_bytes)
                if not formato or not formato[0]:
                    raise Exception("Error al generar el formato de entrega")
        except Exception as e:
            st.error(f"Error al procesar la plantilla: {str(e)}")
            return

        # Verificar tipo de programa
        es_tecnologo = verificar_ficha_tecnologo(df, fichaPdf) == 'Tecnólogo'

        # Definir documentos requeridos
        documentos_base = [
            "F-023(final)",
            "Agencia Publica de Empleo",
            "Paz y Salvo Academico",
            "Copia del Documento de Identidad",
            "Certificación empresa"
        ]
        
        documentos = documentos_base.copy()
        if es_tecnologo:
            documentos.append("Certificación Pruebas TyT")

        # Gestionar archivos subidos
        archivos_subidos = {}
        estado_archivos = st.empty()

        for nombre in documentos:
            archivo = st.file_uploader(nombre, type=["pdf"], key=f"upload_{nombre}")
            if archivo is not None:
                # Leer el contenido del archivo en memoria
                archivo_bytes = BytesIO(archivo.read())
                archivo_bytes.seek(0)
                archivos_subidos[nombre] = archivo_bytes

        # Verificar documentos faltantes
        documentos_faltantes = [doc for doc in documentos if doc not in archivos_subidos]
        
        if documentos_faltantes:
            estado_archivos.warning(f"Documentos obligatorios faltantes: {', '.join(documentos_faltantes)}")
        else:
            estado_archivos.success("Todos los documentos obligatorios han sido subidos")

        # Preparar archivos para unión
        if st.button('Consolidar PDFs', key='btn_consolidar'):
            if documentos_faltantes:
                st.error("Debe subir todos los documentos obligatorios antes de consolidar")
                return

            try:
                with st.spinner('Uniendo PDFs...'):
                    # Crear merger de PDFs
                    merger = PdfMerger()
                    
                    # Añadir resto de documentos en orden
                    for nombre in documentos:
                        if nombre in archivos_subidos:
                            archivos_subidos[nombre].seek(0)
                            merger.append(archivos_subidos[nombre])
                            
                    # Añadir formato generado
                    if formato and formato[0]:
                        formato[0].seek(0)
                        merger.append(formato[0])

                    # Unir PDFs
                    output = BytesIO()
                    merger.write(output)
                    merger.close()
                    output.seek(0)
                    
                    # Verificar el PDF resultante
                    if output.getbuffer().nbytes < 1000:  # menos de 1KB probablemente está corrupto
                        raise Exception("El PDF generado parece estar corrupto")

                    # Preparar nombre del archivo
                    nombre_archivo = f"{fichaPdf} {nombrePdf}.pdf"
                    
                    # Enviar por correo
                    enviar_correo_instructor(
                        f'Consolidado Aprendiz {nombrePdf}',
                        f'El aprendiz {nombrePdf} ha subido los documentos requeridos para la finalización de la etapa productiva.',
                        output,
                        nombre_archivo
                    )

                    # Ofrecer descarga
                    st.download_button(
                        label="Descargar PDF Consolidado",
                        data=output.getvalue(),
                        file_name=nombre_archivo,
                        mime="application/pdf"
                    )

                st.success('El archivo consolidado se ha enviado al instructor exitosamente')
                
            except Exception as e:
                st.error(f'Error al procesar los PDFs: {str(e)}')
                import traceback
                st.error(traceback.format_exc())  # Para debugging
    
    except Exception as e:
        st.error(f"Error general en el formulario: {str(e)}")
        import traceback
        st.error(traceback.format_exc())

def validar_pdf(archivo):
    try:
        if archivo is not None:
            # Verificar el tipo de archivo
            if not archivo.type == "application/pdf":
                return False
            # Intentar leer el PDF para verificar que no está corrupto
            PdfMerger().append(io.BytesIO(archivo.getvalue()))
            return True
    except:
        return False
    return False

# Lógica principal de la aplicación
def main():
    if st.session_state.vista_actual == 'inicio':
        mostrar_inicio()
    elif st.session_state.vista_actual == 'instructor':
        mostrar_instructor()
    elif st.session_state.vista_actual == 'formulario':
        mostrar_formulario()
    elif st.session_state.vista_actual == 'aprendiz':
        mostrar_aprendiz()

# Iniciar la aplicación, de esta forma se evita que si se importa en otro programa no se inicie directamente
if __name__ == "__main__":
    main()
